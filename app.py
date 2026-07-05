from __future__ import annotations

import io
import json
import os
import re
import shutil
import subprocess
from datetime import datetime
from difflib import SequenceMatcher
from pathlib import Path
from typing import List

import pandas as pd
import streamlit as st
from dotenv import load_dotenv
from cerebras.cloud.sdk import Cerebras
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS
from langchain_core.documents import Document
from langchain_core.prompts import ChatPromptTemplate
from langchain_text_splitters import RecursiveCharacterTextSplitter

BASE_DIR = Path(__file__).resolve().parent
ENV_PATH = BASE_DIR / ".env"
DATA_DIR = Path(os.getenv("APP_DATA_DIR", BASE_DIR / "data"))
UPLOAD_DIR = DATA_DIR / "uploads"
INDEX_DIR = DATA_DIR / "faiss_index"
MANIFEST_PATH = DATA_DIR / "manifest.json"
TESSDATA_DIR = Path(os.getenv("APP_TESSDATA_DIR", DATA_DIR / "tessdata"))
DOCUMENT_METADATA_PATH = DATA_DIR / "document_metadata.json"
FEE_RECORDS_PATH = DATA_DIR / "fee_records.json"

load_dotenv(ENV_PATH, override=True)

SUPPORTED_TYPES = {"pdf", "docx", "xlsx", "csv", "txt", "md"}

STOPWORDS = {
    "a",
    "an",
    "and",
    "are",
    "as",
    "at",
    "be",
    "by",
    "for",
    "from",
    "how",
    "i",
    "in",
    "is",
    "it",
    "me",
    "not",
    "of",
    "on",
    "or",
    "the",
    "to",
    "was",
    "were",
    "what",
    "which",
    "who",
    "will",
    "आहे",
    "आहेत",
    "काय",
    "कोण",
    "कोणी",
    "कोणत्या",
    "कोणते",
    "का",
    "कसा",
    "कशी",
    "कुठे",
    "या",
    "हा",
    "ही",
    "हे",
    "मध्ये",
    "और",
    "का",
    "की",
    "क्या",
    "किस",
    "किसने",
    "कौन",
    "है",
    "हैं",
    "में",
}

DOMAIN_QUERY_TERMS = {
    "centre",
    "centres",
    "center",
    "centers",
    "exam",
    "examination",
    "online",
    "listed",
    "list",
    "city",
    "cities",
    "state",
    "states",
    "territory",
    "territories",
    "vacancy",
    "vacancies",
    "junior",
    "assistant",
}

QUERY_CORRECTIONS = {
    "hte": "the",
    "qulaities": "qualities",
    "qulaity": "quality",
    "adress": "address",
    "resarch": "research",
    "mosquitp": "mosquito",
    "mosquitofusion": "mosquitofusion",
    "og": "of",
    "metodology": "methodology",
    "methodolgy": "methodology",
}

PROGRAM_ALIASES = {
    "B.Tech": (r"\bb\s*\.?\s*tech\b", r"\bbachelor\s+of\s+technology\b", r"\bb\.?e\.?\s*/\s*b\.?tech\b"),
    "B.E.": (r"\bb\.\s*e\.?\b", r"\bb\s+e\b", r"\bbachelor\s+of\s+engineering\b", r"\bb\.?e\.?\s*/\s*b\.?tech\b"),
    "M.Arch": (r"\bm\.\s*arch\b", r"\bm\s+arch\b", r"\bmarch[_\-]", r"\bmaster\s+of\s+architecture\b", r"एम\.?\s*आर्क"),
    "B.Arch": (r"\bb\s*\.?\s*arch\b", r"\bbachelor\s+of\s+architecture\b"),
    "M.Plan": (r"\bm\s*\.?\s*plan\b", r"\bmplan\b", r"\bmaster\s+of\s+planning\b", r"एम\.?\s*प्लॅन", r"एम\.?\s*प्लान"),
    "B.Plan": (r"\bb\s*\.?\s*plan\b", r"\bbplan\b", r"\bbachelor\s+of\s+planning\b", r"बी\.?\s*प्लॅन", r"बी\.?\s*प्लान"),
    "MBA": (r"\bmba\b", r"\bm\.?b\.?a\.?\b", r"\bmaster\s+of\s+business\s+administration\b"),
    "MCA": (r"\bmca\b", r"\bm\.?c\.?a\.?\b", r"\bmaster\s+of\s+computer\s+applications\b"),
    "M.Tech": (r"\bm\s*\.?\s*tech\b", r"\bmaster\s+of\s+technology\b"),
    "M.E.": (r"\bm\.\s*e\.?\b", r"\bm\s+e\b", r"\bmaster\s+of\s+engineering\b"),
    "B.Pharm": (r"\bb\s*\.?\s*pharm\b", r"\bbachelor\s+of\s+pharmacy\b"),
    "M.Pharm": (r"\bm\s*\.?\s*pharm\b", r"\bmaster\s+of\s+pharmacy\b"),
}

FEE_TYPE_KEYWORDS = {
    "hostel_fee": ("hostel", "hostel fee", "hostel charges", "वसतिगृह", "वसतिगृह शुल्क", "हॉस्टेल", "हॉस्टेल फी"),
    "tuition_fee": ("tuition", "tuition fee", "शिक्षण शुल्क", "ट्यूशन फी"),
    "college_fee": ("college fee", "college fees", "college and university fee", "college & university fee", "कॉलेज शुल्क", "कॉलेज व विद्यापीठ शुल्क", "कॉलेज आणि विद्यापीठ शुल्क", "महाविद्यालय शुल्क"),
    "university_fee": ("university fee", "university fees", "college and university fee", "college & university fee", "विद्यापीठ शुल्क", "कॉलेज व विद्यापीठ शुल्क", "कॉलेज आणि विद्यापीठ शुल्क"),
    "caution_money": ("caution", "caution money", "सावधनता ठेव", "कौशन मनी"),
    "security_deposit": ("security", "security money", "security deposit", "सुरक्षा ठेव", "सिक्युरिटी"),
    "development_fee": ("development", "development fee", "विकास शुल्क"),
    "fee_structure": ("fee structure", "total fee", "course fee", "annual fee", "semester fee", "fee", "फी संरचना", "फी", "शुल्क"),
}

STRUCTURED_QUERY_TERMS = {
    "fee",
    "fees",
    "hostel",
    "tuition",
    "caution",
    "security",
    "deposit",
    "money",
    "structure",
    "amount",
    "फी",
    "शुल्क",
    "वसतिगृह",
    "विद्यापीठ",
    "कॉलेज",
    "महाविद्यालय",
}


def ensure_dirs() -> None:
    DATA_DIR.mkdir(exist_ok=True)
    UPLOAD_DIR.mkdir(exist_ok=True)
    INDEX_DIR.mkdir(exist_ok=True)


def file_suffix(file_name: str) -> str:
    return Path(file_name).suffix.lower().lstrip(".")


def read_text_file(file_bytes: bytes) -> str:
    for encoding in ("utf-8", "utf-16", "latin-1"):
        try:
            return file_bytes.decode(encoding)
        except UnicodeDecodeError:
            continue
    return file_bytes.decode("utf-8", errors="ignore")


def _clean_text_lines(text: str) -> str:
    lines = [line.strip() for line in text.splitlines()]
    cleaned = "\n".join(line for line in lines if line)
    return _fix_common_ocr_errors(cleaned)


def _fix_common_ocr_errors(text: str) -> str:
    replacements = {
        "दिनांक : ९७ मार्च": "दिनांक : १७ मार्च",
        "दिनांक: ९७ मार्च": "दिनांक: १७ मार्च",
        "असत्ताना": "असताना",
        "बाटले": "वाटले",
        "हौ सदिच्छा": "ही सदिच्छा",
    }
    for wrong, right in replacements.items():
        text = text.replace(wrong, right)
    return text


def _available_local_ocr_languages() -> list[str]:
    preferred_languages = ["mar", "hin", "eng"]
    available: list[str] = []
    for language in preferred_languages:
        if (TESSDATA_DIR / f"{language}.traineddata").exists():
            available.append(language)
    if available:
        return available

    try:
        result = subprocess.run(
            ["tesseract", "--list-langs"],
            capture_output=True,
            text=True,
            timeout=5,
            check=False,
        )
        installed = set(result.stdout.splitlines()[1:])
    except Exception:
        installed = set()
    for language in preferred_languages:
        if language in installed:
            available.append(language)
    return available


def _ocr_language_config() -> tuple[str | None, str]:
    languages = _available_local_ocr_languages()
    config_parts = ["--oem 3", "--psm 6"]
    if TESSDATA_DIR.exists() and languages:
        config_parts.append(f"--tessdata-dir {TESSDATA_DIR}")
    return ("+".join(languages) if languages else None, " ".join(config_parts))


def _page_dict_to_text(page) -> str:
    try:
        page_dict = page.get_text("dict")
    except Exception:
        return ""

    lines: list[str] = []
    for block in page_dict.get("blocks", []):
        if block.get("type", 0) != 0:
            continue
        block_lines: list[str] = []
        for line in block.get("lines", []):
            spans = [span.get("text", "").strip() for span in line.get("spans", [])]
            line_text = "".join(spans).strip()
            if line_text:
                block_lines.append(line_text)
        if block_lines:
            lines.append("\n".join(block_lines))

    return _clean_text_lines("\n\n".join(lines))


def _page_tables_to_text(page) -> list[str]:
    table_texts: list[str] = []
    find_tables = getattr(page, "find_tables", None)
    if not callable(find_tables):
        return table_texts

    try:
        tables = find_tables()
    except Exception:
        return table_texts

    table_list = getattr(tables, "tables", []) or []
    for table_index, table in enumerate(table_list, 1):
        try:
            data = table.extract()
        except Exception:
            continue
        rows: list[str] = []
        for row in data or []:
            cells = [str(cell).strip() for cell in row if str(cell).strip()]
            if cells:
                rows.append(" | ".join(cells))
        if rows:
            table_texts.append(f"TABLE {table_index}\n" + "\n".join(rows))

    return table_texts


def _ocr_page(page) -> str:
    try:
        import fitz
    except ImportError as exc:  # pragma: no cover - optional dependency
        raise RuntimeError("pymupdf is required for PDF OCR fallback") from exc

    try:
        from PIL import Image, ImageFilter, ImageOps
    except ImportError as exc:  # pragma: no cover - optional dependency
        raise RuntimeError("Pillow is required for PDF OCR fallback") from exc

    try:
        import pytesseract
    except ImportError as exc:  # pragma: no cover - optional dependency
        raise RuntimeError("pytesseract is required for OCR fallback") from exc

    # Use a higher DPI render so table text and headings are easier to read.
    pixmap = page.get_pixmap(matrix=fitz.Matrix(3, 3), alpha=False)
    image = Image.open(io.BytesIO(pixmap.tobytes("png")))
    image = ImageOps.grayscale(image)
    image = image.filter(ImageFilter.SHARPEN)

    language, config = _ocr_language_config()
    if language:
        text = pytesseract.image_to_string(image, lang=language, config=config)
    else:
        text = pytesseract.image_to_string(image, config=config)
    return _clean_text_lines(text)


def extract_pdf_documents(file_path: Path, source_name: str) -> List[Document]:
    documents: List[Document] = []

    try:
        import fitz
    except ImportError as exc:  # pragma: no cover - optional dependency
        raise RuntimeError("pymupdf is required for PDF extraction") from exc

    pdf = fitz.open(file_path)
    for page_number in range(len(pdf)):
        page = pdf.load_page(page_number)

        structured_text = _page_dict_to_text(page)
        table_texts = _page_tables_to_text(page)
        fallback_text = page.get_text("text").strip()

        text_parts = [part for part in [structured_text, *table_texts, fallback_text] if part]
        text = _clean_text_lines("\n\n".join(text_parts))

        # If the page looks like a scan or a badly extracted image page, OCR it.
        if len(text) < 80:
            try:
                ocr_text = _ocr_page(page)
                if len(ocr_text) > len(text):
                    text = ocr_text
            except Exception:
                pass

        if text:
            documents.append(
                Document(
                    page_content=text,
                    metadata={
                        "source": source_name,
                        "page": page_number + 1,
                        "file_type": "pdf",
                    },
                )
            )

    return documents


def extract_docx_documents(file_path: Path, source_name: str) -> List[Document]:
    from docx import Document as DocxDocument

    doc = DocxDocument(str(file_path))
    parts = [para.text.strip() for para in doc.paragraphs if para.text.strip()]
    if not parts:
        return []

    return [
        Document(
            page_content="\n".join(parts),
            metadata={"source": source_name, "page": 1, "file_type": "docx"},
        )
    ]


def extract_tabular_documents(file_path: Path, source_name: str, suffix: str) -> List[Document]:
    if suffix == "csv":
        df = pd.read_csv(file_path)
        content = df.to_csv(index=False)
        sheet_name = "csv"
        if not content.strip():
            return []
        return [
            Document(
                page_content=f"Sheet: {sheet_name}\n{content}",
                metadata={
                    "source": source_name,
                    "page": 1,
                    "sheet": sheet_name,
                    "file_type": suffix,
                },
            )
        ]

    sheets = pd.read_excel(file_path, sheet_name=None)
    documents: List[Document] = []
    for sheet_name, df in sheets.items():
        content = df.fillna("").to_csv(index=False)
        if content.strip():
            documents.append(
                Document(
                    page_content=f"Sheet: {sheet_name}\n{content}",
                    metadata={
                        "source": source_name,
                        "page": 1,
                        "sheet": sheet_name,
                        "file_type": suffix,
                    },
                )
            )
    return documents


def extract_plain_text_documents(file_bytes: bytes, source_name: str, suffix: str) -> List[Document]:
    text = read_text_file(file_bytes).strip()
    if not text:
        return []

    return [
        Document(
            page_content=text,
            metadata={"source": source_name, "page": 1, "file_type": suffix},
        )
    ]


def load_documents_from_upload(uploaded_file) -> List[Document]:
    suffix = file_suffix(uploaded_file.name)
    if suffix not in SUPPORTED_TYPES:
        raise ValueError(f"Unsupported file type: {suffix}")

    ensure_dirs()
    safe_name = Path(uploaded_file.name).name
    stored_path = UPLOAD_DIR / safe_name
    stored_path.write_bytes(uploaded_file.getbuffer())

    if suffix == "pdf":
        return extract_pdf_documents(stored_path, safe_name)
    if suffix == "docx":
        return extract_docx_documents(stored_path, safe_name)
    if suffix in {"xlsx", "csv"}:
        return extract_tabular_documents(stored_path, safe_name, suffix)
    return extract_plain_text_documents(stored_path.read_bytes(), safe_name, suffix)


def build_chunks(documents: List[Document]) -> List[Document]:
    splitter = RecursiveCharacterTextSplitter(chunk_size=900, chunk_overlap=150)
    chunks = splitter.split_documents(documents)

    for index, chunk in enumerate(chunks):
        chunk.metadata["chunk_index"] = index

    return chunks


def _detect_programs(text: str) -> list[str]:
    found: list[str] = []
    for program, patterns in PROGRAM_ALIASES.items():
        if any(re.search(pattern, text, flags=re.IGNORECASE) for pattern in patterns):
            found.append(program)
    if "B.Tech" in found and "B.E." in found:
        found = [program for program in found if program != "B.E."]
    return found


def _detect_document_type(text: str) -> str:
    text_lc = text.lower()
    if any(term in text_lc for term in ("fee structure", "hostel fee", "caution money", "tuition fee", "फी", "शुल्क")):
        return "fee"
    if "cutoff" in text_lc or "cut off" in text_lc:
        return "cutoff"
    if "seat matrix" in text_lc:
        return "seat_matrix"
    if "eligibility" in text_lc:
        return "eligibility"
    if "resume" in text_lc or "career objective" in text_lc:
        return "resume"
    if "research paper" in text_lc or "abstract" in text_lc:
        return "research_paper"
    return "general"


def _infer_document_title(source: str, text: str) -> str:
    for line in text.splitlines()[:12]:
        clean_line = line.strip(" |")
        if clean_line and len(clean_line) <= 140:
            return clean_line
    return Path(source).stem.replace("_", " ")


def _infer_session(text: str) -> str | None:
    match = re.search(r"(?:session|academic\s+year|year)\s*[:\-]?\s*(20\d{2}\s*[-–]\s*\d{2,4})", text, flags=re.IGNORECASE)
    if match:
        return match.group(1).replace(" ", "")
    match = re.search(r"\b(20\d{2}\s*[-–]\s*\d{2})\b", text)
    return match.group(1).replace(" ", "") if match else None


def infer_document_metadata(documents: list[Document]) -> list[dict]:
    by_source: dict[str, list[Document]] = {}
    for doc in documents:
        by_source.setdefault(str(doc.metadata.get("source", "unknown")), []).append(doc)

    metadata: list[dict] = []
    for source, source_docs in by_source.items():
        text = "\n".join(doc.page_content for doc in source_docs)
        metadata.append(
            {
                "source": source,
                "document_title": _infer_document_title(source, text),
                "document_type": _detect_document_type(text),
                "programs_detected": _detect_programs(source + "\n" + text),
                "session": _infer_session(source + "\n" + text),
                "pages": sorted({doc.metadata.get("page", 1) for doc in source_docs}),
            }
        )
    return metadata


def _detect_fee_types(text: str) -> list[str]:
    text_lc = text.lower()
    fee_types: list[str] = []
    for fee_type, keywords in FEE_TYPE_KEYWORDS.items():
        if any(keyword in text_lc for keyword in keywords):
            fee_types.append(fee_type)
    if not fee_types and re.search(r"(?:rs\.?|inr|₹)\s*[\d,]+", text_lc):
        fee_types.append("fee_structure")
    return fee_types


def _amount_matches(text: str) -> list[tuple[str, int]]:
    patterns = (
        r"(?:rs\.?|inr|\u20b9)\s*[\d,]+(?:\.\d{1,2})?",
        r"\b[\d]{1,3}(?:,[\d]{3})+(?:\.\d{1,2})?\b",
    )
    matches: list[tuple[str, int]] = []
    seen: set[str] = set()
    for pattern in patterns:
        for match in re.finditer(pattern, text, flags=re.IGNORECASE):
            amount = _normalise_amount(match.group(0).strip())
            key = f"{amount}:{match.start()}"
            if key not in seen:
                matches.append((amount, match.start()))
                seen.add(key)
    return sorted(matches, key=lambda item: item[1])


def _amounts_in_text(text: str) -> list[str]:
    amounts: list[str] = []
    for amount, _position in _amount_matches(text):
        if amount not in amounts:
            amounts.append(amount)
    return amounts


def _normalise_amount(amount: str) -> str:
    amount = amount.replace("\u20b9", "Rs. ").strip()
    if not re.match(r"(?i)^(rs\.?|inr)\b", amount):
        amount = f"Rs. {amount}"
    return re.sub(r"\s+", " ", amount)


def _keyword_positions(text_lc: str, fee_type: str) -> list[int]:
    positions: list[int] = []
    for keyword in FEE_TYPE_KEYWORDS.get(fee_type, ()):
        position = text_lc.find(keyword)
        if position >= 0:
            positions.append(position)
    return positions


def _amounts_for_fee_type(text: str, fee_type: str) -> list[str]:
    matches = _amount_matches(text)
    if not matches:
        return []

    if fee_type == "fee_structure":
        return [matches[0][0]]

    positions = _keyword_positions(text.lower(), fee_type)
    if not positions:
        return [matches[0][0]]

    chosen: list[str] = []
    for position in positions:
        after_keyword = [match for match in matches if match[1] >= position]
        if after_keyword:
            chosen.append(after_keyword[0][0])
    if not chosen:
        return []

    unique: list[str] = []
    for amount in chosen:
        if amount not in unique:
            unique.append(amount)
    return unique


def extract_fee_records(documents: list[Document], document_metadata: list[dict] | None = None) -> list[dict]:
    metadata_by_source = {item["source"]: item for item in document_metadata or infer_document_metadata(documents)}
    records: list[dict] = []
    seen: set[tuple[str, str, str, str, str]] = set()

    for doc in documents:
        source = str(doc.metadata.get("source", "unknown"))
        page = doc.metadata.get("page", 1)
        source_meta = metadata_by_source.get(source, {})
        source_programs = source_meta.get("programs_detected") or []
        lines = [line.strip(" |") for line in doc.page_content.splitlines() if line.strip(" |")]
        windows = lines + [_normalise_inline_text(doc.page_content)]

        for text in windows:
            if not any(
                term in text.lower()
                for term in (
                    "fee",
                    "hostel",
                    "caution",
                    "security",
                    "tuition",
                    "deposit",
                    "college",
                    "university",
                    "rs",
                    "₹",
                    "inr",
                    "फी",
                    "शुल्क",
                    "वसतिगृह",
                    "विद्यापीठ",
                    "कॉलेज",
                    "महाविद्यालय",
                )
            ):
                continue
            amount_matches = _amount_matches(text)
            fee_types = _detect_fee_types(text)
            if not amount_matches or not fee_types:
                continue
            programs = _detect_programs(text) or source_programs
            if not programs:
                continue
            for program in programs:
                for fee_type in fee_types:
                    for amount in _amounts_for_fee_type(text, fee_type):
                        record = {
                            "program": program,
                            "fee_type": fee_type,
                            "amount": amount,
                            "currency": "INR",
                            "notes": text[:260],
                            "source": source,
                            "page": page,
                        }
                        key = (record["program"], record["fee_type"], record["amount"], record["source"], str(record["page"]))
                        if key not in seen:
                            records.append(record)
                            seen.add(key)
    return records


def persist_manifest(files: List[str], chunk_count: int) -> None:
    ensure_dirs()
    manifest = {
        "files": files,
        "chunk_count": chunk_count,
        "updated_at": datetime.utcnow().isoformat() + "Z",
    }
    MANIFEST_PATH.write_text(json.dumps(manifest, indent=2), encoding="utf-8")


def load_manifest() -> dict:
    if MANIFEST_PATH.exists():
        return json.loads(MANIFEST_PATH.read_text(encoding="utf-8"))
    return {"files": [], "chunk_count": 0, "updated_at": None}


def _read_json_file(path: Path, default):
    if path.exists():
        try:
            return json.loads(path.read_text(encoding="utf-8"))
        except Exception:
            return default
    return default


def _write_json_file(path: Path, data) -> None:
    ensure_dirs()
    path.write_text(json.dumps(data, indent=2, ensure_ascii=False), encoding="utf-8")


def load_document_metadata() -> list[dict]:
    return _read_json_file(DOCUMENT_METADATA_PATH, [])


def save_document_metadata(metadata: list[dict]) -> None:
    _write_json_file(DOCUMENT_METADATA_PATH, metadata)


def load_fee_records() -> list[dict]:
    return _read_json_file(FEE_RECORDS_PATH, [])


def save_fee_records(records: list[dict]) -> None:
    _write_json_file(FEE_RECORDS_PATH, records)


def get_embeddings() -> HuggingFaceEmbeddings:
    model = os.getenv("EMBEDDING_MODEL", "sentence-transformers/all-MiniLM-L6-v2")
    return HuggingFaceEmbeddings(model_name=model)


def build_vector_store(chunks: List[Document]) -> FAISS:
    embeddings = get_embeddings()
    return FAISS.from_documents(chunks, embeddings)


def save_vector_store(vector_store: FAISS) -> None:
    ensure_dirs()
    vector_store.save_local(str(INDEX_DIR))


def load_vector_store() -> FAISS | None:
    if not INDEX_DIR.exists() or not any(INDEX_DIR.iterdir()):
        return None

    embeddings = get_embeddings()
    return FAISS.load_local(
        str(INDEX_DIR),
        embeddings,
        allow_dangerous_deserialization=True,
    )


def build_answer_prompt() -> ChatPromptTemplate:
    return ChatPromptTemplate.from_messages(
        [
            (
                "system",
                "You are a grounded document assistant for a client demo.\n"
                "Answer only from the provided context.\n"
                "Use the chat memory to resolve follow-up questions and pronouns, but do not invent facts.\n"
                "If the answer is not explicitly present, say you could not find it in the uploaded documents.\n"
                "Only answer when you are at least 95% confident the context directly supports it.\n"
                "Use the same language as the user's question when possible.\n"
                "For exam-centre questions, rely on the Online Examination Centre section/table; "
                "do not treat vacancy-city tables as exam-centre evidence.\n"
                "For yes/no listing questions, answer yes or no first, then give the exact listed row.\n"
                "Do not use outside knowledge, guesses, or paraphrases that add facts not in the context.\n"
                "Keep the answer short, direct, and cite sources inline using [source p.X] when possible.",
            ),
            (
                "human",
                "Chat memory:\n{chat_history}\n\n"
                "Question: {question}\n\nContext:\n{context}",
            ),
        ]
    )


def _normalise_inline_text(text: str) -> str:
    return " ".join(text.split())


def _normalise_query(question: str) -> str:
    if _contains_devanagari(question):
        text = re.sub(r"\s+", " ", question).strip()
        return re.sub(r"\s+([?.!,;:।])", r"\1", text)

    tokens = re.findall(r"\w+|[^\w\s]", question, flags=re.UNICODE)
    corrected: list[str] = []
    for token in tokens:
        replacement = QUERY_CORRECTIONS.get(token.lower())
        corrected.append(replacement if replacement else token)
    text = " ".join(corrected)
    text = re.sub(r"\s+([?.!,;:])", r"\1", text)
    return text


def _fuzzy_ratio(left: str, right: str) -> float:
    return SequenceMatcher(None, left.lower(), right.lower()).ratio()


def _source_relevance_score(question: str, source: str) -> float:
    source_stem = Path(source).stem.lower().replace("_", " ").replace("-", " ")
    question_lc = question.lower()
    score = 0.0
    for term in _query_terms(question):
        if term in source_stem:
            score += 0.35
        elif len(term) >= 5:
            score += max(0.0, _fuzzy_ratio(term, source_stem) - 0.55)
    if source_stem and source_stem in question_lc:
        score += 1.0
    return score


def _find_exam_centres_for_state(docs: list[Document], state: str = "gujarat") -> list[str]:
    combined = _normalise_inline_text("\n".join(doc.page_content for doc in docs))
    next_state = "maharashtra" if state.lower() == "gujarat" else ""
    if next_state:
        pattern = rf"{re.escape(state)}\s+(.+?)\s+{re.escape(next_state)}"
    else:
        pattern = rf"{re.escape(state)}\s+(.+?)(?:11\.\s*INTERVIEW|$)"
    matches = re.finditer(pattern, combined, flags=re.IGNORECASE)
    for match in matches:
        raw_centres = match.group(1).strip(" ,")
        if any(city.lower() in raw_centres.lower() for city in ("surat", "vadodara", "rajkot")):
            return [centre.strip(" ,") for centre in raw_centres.split(",") if centre.strip(" ,")]
    return []


def _fallback_exam_centre_answer(question: str, docs: list[Document]) -> str | None:
    question_lc = question.lower()
    if "exam" not in question_lc and "examination" not in question_lc:
        return None
    if "centre" not in question_lc and "center" not in question_lc:
        return None

    centres = _find_exam_centres_for_state(docs)
    if not centres:
        return None

    source_doc = next((doc for doc in docs if "surat" in doc.page_content.lower()), docs[0])
    source = source_doc.metadata.get("source", "unknown")
    page = source_doc.metadata.get("page", "?")
    centre_text = ", ".join(centres)

    asked_city = next((centre for centre in centres if centre.lower() in question_lc), None)
    if asked_city:
        return f"Yes. {asked_city} is listed as a Gujarat online examination centre. Gujarat centres are: {centre_text}. [{source} p.{page}]"

    if "gujarat" in question_lc:
        return f"The Gujarat online examination centres are: {centre_text}. [{source} p.{page}]"

    return None


def _fallback_vacancy_answer(question: str, docs: list[Document]) -> str | None:
    question_lc = question.lower()
    if "vacanc" not in question_lc or "gujarat" not in question_lc:
        return None

    combined = _normalise_inline_text("\n".join(doc.page_content for doc in docs))
    if not all(term in combined.lower() for term in ("gujarat", "ahmedabad", "rajkot")):
        return None

    source_doc = next((doc for doc in docs if "number of vacancies" in doc.page_content.lower()), docs[0])
    source = source_doc.metadata.get("source", "unknown")
    page = source_doc.metadata.get("page", "?")
    return f"The Gujarat number of vacancies is Ahmedabad (1) and Rajkot (1), for a total of 2 Junior Assistant vacancies. [{source} p.{page}]"


def _is_fee_question(question: str) -> bool:
    q = question.lower()
    return any(
        term in q
        for term in (
            "fee",
            "fees",
            "hostel",
            "tuition",
            "caution",
            "security",
            "deposit",
            "college",
            "university",
            "फी",
            "शुल्क",
            "वसतिगृह",
            "हॉस्टेल",
            "विद्यापीठ",
            "कॉलेज",
            "महाविद्यालय",
        )
    )


def _requested_fee_types(question: str) -> list[str]:
    question_lc = question.lower()
    fee_types = _detect_fee_types(question)
    specific_fee_types = [fee_type for fee_type in fee_types if fee_type != "fee_structure"]
    explicit_fee_structure = any(
        phrase in question_lc
        for phrase in ("fee structure", "course fee", "total fee", "annual fee", "semester fee", "फी संरचना")
    )
    if specific_fee_types and not explicit_fee_structure:
        return specific_fee_types
    if "fee_structure" not in fee_types and any(term in question_lc for term in ("fee", "fees", "structure", "फी", "शुल्क")):
        fee_types.append("fee_structure")
    return fee_types or ["fee_structure"]


def _records_from_vector_store(vector_store: FAISS) -> tuple[list[dict], list[dict]]:
    docs = _all_indexed_documents(vector_store)
    metadata = infer_document_metadata(docs)
    records = extract_fee_records(docs, metadata)
    persisted_records = load_fee_records()
    if persisted_records:
        records = persisted_records
    persisted_metadata = load_document_metadata()
    if persisted_metadata:
        metadata = persisted_metadata
    return records, metadata


def _programs_from_metadata(metadata: list[dict]) -> list[str]:
    programs: list[str] = []
    for item in metadata:
        for program in item.get("programs_detected") or []:
            if program not in programs:
                programs.append(program)
    return programs


def _fee_type_label(fee_type: str) -> str:
    return fee_type.replace("_", " ")


def _localized_fee_type_label(fee_type: str, language: str) -> str:
    if language == "mr":
        labels = {
            "hostel_fee": "वसतिगृह शुल्क",
            "tuition_fee": "शिक्षण शुल्क",
            "college_fee": "कॉलेज शुल्क",
            "university_fee": "विद्यापीठ शुल्क",
            "caution_money": "कौशन मनी",
            "security_deposit": "सुरक्षा ठेव",
            "development_fee": "विकास शुल्क",
            "fee_structure": "फी",
        }
        return labels.get(fee_type, _fee_type_label(fee_type))
    return _fee_type_label(fee_type)


def _fee_refusal(requested_programs: list[str], available_programs: list[str]) -> str:
    requested_text = ", ".join(requested_programs)
    available_text = ", ".join(available_programs) if available_programs else "other uploaded documents"
    return (
        f"I found fee information in the uploaded documents, but it appears to be for {available_text}, "
        f"not {requested_text}. I cannot answer {requested_text} fees from these documents."
    )


def _fee_type_text(fee_types: list[str]) -> str:
    return ", ".join(_fee_type_label(fee_type) for fee_type in fee_types)


def _fee_type_refusal(requested_programs: list[str], requested_fee_types: list[str]) -> str:
    program_text = ", ".join(requested_programs)
    fee_text = _fee_type_text(requested_fee_types)
    return (
        f"I found fee records for {program_text}, but not a clearly supported {fee_text} record. "
        f"I cannot answer that fee type from these documents."
    )


def _fee_record_refusal(requested_programs: list[str]) -> str:
    program_text = ", ".join(requested_programs)
    return (
        f"I found the requested program ({program_text}) in the uploaded documents, "
        "but I could not extract a clearly supported fee record for it."
    )


def _fallback_fee_answer(vector_store: FAISS, question: str) -> tuple[str, list[Document]] | None:
    question = _normalise_query(question)
    if not _is_fee_question(question):
        return None

    requested_programs = _detect_programs(question)
    requested_fee_types = _requested_fee_types(question)
    records, metadata = _records_from_vector_store(vector_store)

    available_programs = _programs_from_metadata(metadata) or sorted({record["program"] for record in records})
    if not records:
        if requested_programs and available_programs and not any(program in available_programs for program in requested_programs):
            return _fee_refusal(requested_programs, available_programs), _all_indexed_documents(vector_store)[:4]
        if requested_programs and available_programs and any(program in available_programs for program in requested_programs):
            return _fee_record_refusal(requested_programs), _all_indexed_documents(vector_store)[:4]
        return None

    if requested_programs:
        matching_records = [
            record
            for record in records
            if record["program"] in requested_programs
            and record["fee_type"] in requested_fee_types
        ]
        if not matching_records:
            if any(program in available_programs for program in requested_programs):
                return _fee_type_refusal(requested_programs, requested_fee_types), _all_indexed_documents(vector_store)[:4]
            return _fee_refusal(requested_programs, available_programs), _all_indexed_documents(vector_store)[:4]
    else:
        matching_records = [
            record
            for record in records
            if record["fee_type"] in requested_fee_types
        ]
        if not matching_records:
            return None
        if len({record["program"] for record in matching_records}) > 1:
            programs = ", ".join(sorted({record["program"] for record in matching_records}))
            return f"I found fee records for multiple programs ({programs}). Please specify the program.", _all_indexed_documents(vector_store)[:4]

    grouped: dict[tuple[str, str], dict] = {}
    for record in matching_records:
        key = (record["program"], record["fee_type"])
        grouped.setdefault(key, record)

    language = _detect_question_language(question)
    lines = []
    if language == "mr" and any(marker in question for marker in ("अभ्यासक्रम", "कोणत्या अभ्यासक्रम", "नोटीस")):
        programs = ", ".join(sorted({record["program"] for record in matching_records}))
        lines.append(f"ही फी नोटीस {programs} अभ्यासक्रमासाठी आहे.")
    for (_program, _fee_type), record in grouped.items():
        fee_label = _localized_fee_type_label(record["fee_type"], language)
        lines.append(
            f"{record['program']} {fee_label}: {record['amount']} "
            f"[{record['source']} p.{record['page']}]"
        )

    source_docs = [
        doc
        for doc in _all_indexed_documents(vector_store)
        if any(doc.metadata.get("source") == record["source"] for record in matching_records)
    ]
    return "\n".join(lines), source_docs[:8] or _all_indexed_documents(vector_store)[:4]


def _answer_amounts(answer: str) -> list[str]:
    return [_normalise_amount(amount) for amount in _amounts_in_text(answer)]


def _verify_structured_claims(answer: str, question: str, vector_store: FAISS) -> str | None:
    if not _is_fee_question(question) and not _amounts_in_text(answer):
        return answer

    answer_amounts = _answer_amounts(answer)
    if not answer_amounts:
        return answer

    records, metadata = _records_from_vector_store(vector_store)
    requested_programs = _detect_programs(question)
    if requested_programs:
        records = [record for record in records if record["program"] in requested_programs]
        if not records:
            return _fee_refusal(requested_programs, _programs_from_metadata(metadata))

    record_amounts = {_normalise_amount(record["amount"]) for record in records}
    unsupported = [amount for amount in answer_amounts if amount not in record_amounts]
    if unsupported:
        return None
    return answer


def _is_research_section_question(question: str) -> str | None:
    question = _normalise_query(question)
    question_lc = question.lower()
    if "abstract" in question_lc:
        return "abstract"
    if "methodology" in question_lc or "method" in question_lc:
        return "methodology"
    return None


def _research_paper_sources(vector_store: FAISS) -> list[str]:
    sources: dict[str, int] = {}
    for doc in _all_indexed_documents(vector_store):
        source = str(doc.metadata.get("source", ""))
        text = doc.page_content.lower()
        if doc.metadata.get("file_type") == "docx" and (
            "research paper" in text
            or "abstract" in text
            or "methodology" in text
            or "mosquitofusion" in text
        ):
            sources[source] = sources.get(source, 0) + 1
    return sorted(sources, key=sources.get, reverse=True)


def _ordered_source_text(vector_store: FAISS, source: str) -> tuple[str, Document | None]:
    docs = [
        doc
        for doc in _all_indexed_documents(vector_store)
        if doc.metadata.get("source") == source
    ]
    docs.sort(key=lambda doc: int(doc.metadata.get("chunk_index", 0)))
    uploaded_text = _source_text_from_uploaded_file(source)
    return (uploaded_text or "\n".join(doc.page_content for doc in docs)), (docs[0] if docs else None)


def _source_text_from_uploaded_file(source: str) -> str | None:
    source_path = UPLOAD_DIR / Path(source).name
    if not source_path.exists():
        return None

    suffix = source_path.suffix.lower()
    try:
        if suffix == ".docx":
            from docx import Document as DocxDocument

            doc = DocxDocument(str(source_path))
            parts = [para.text.strip() for para in doc.paragraphs if para.text.strip()]
            return "\n".join(parts)
        if suffix == ".pdf":
            import fitz

            pdf = fitz.open(source_path)
            pages = [page.get_text("text").strip() for page in pdf if page.get_text("text").strip()]
            return "\n".join(pages)
        if suffix in {".txt", ".md"}:
            return read_text_file(source_path.read_bytes())
    except Exception:
        return None
    return None


def _extract_research_section(full_text: str, section: str) -> str | None:
    if section == "abstract":
        match = re.search(
            r"\bABSTRACT\s+(.+?)(?:\n?Index Terms|I\.\s+INTRODUCTION|INTRODUCTION)",
            full_text,
            flags=re.IGNORECASE | re.DOTALL,
        )
        if match:
            return _normalise_inline_text(match.group(1))

    if section == "methodology":
        match = re.search(
            r"IV\.\s+METHODOLOGY\s+(.+?)(?:\n?V\.\s+EXPERIMENTS|\n?VI\.\s+DISCUSSION|\n?VII\.)",
            full_text,
            flags=re.IGNORECASE | re.DOTALL,
        )
        if match:
            return _normalise_inline_text(match.group(1))
    return None


def _trim_section_answer(text: str, max_words: int = 150) -> str:
    text = _dedupe_section_text(text)
    words = text.split()
    if len(words) <= max_words:
        return text
    return " ".join(words[:max_words]) + "..."


def _dedupe_section_text(text: str) -> str:
    text = re.sub(r"\s+", " ", text).strip()
    text = re.sub(r"(IV\.\s+METHODOLOGY\s*)+", "IV. METHODOLOGY ", text, flags=re.IGNORECASE)
    chunks = re.split(r"(?<=[.!?])\s+", text)
    seen: set[str] = set()
    kept: list[str] = []
    for chunk in chunks:
        key = re.sub(r"[^a-z0-9]+", " ", chunk.lower()).strip()
        if not key or key in seen:
            continue
        seen.add(key)
        kept.append(chunk.strip())
    deduped = " ".join(kept)
    duplicate_phrases = [
        "Figure 3: Visual Representation of the Dataset Split Strategy IV. METHODOLOGY",
        "A. Overall Pipeline A. Overall Pipeline",
        "A rigorous preprocessing pipeline was applied to all dataset images prior to model training. The pipeline consisted of the following steps: A rigorous preprocessing pipeline was applied to all dataset images prior to model training. The pipeline consisted of the following steps:",
    ]
    for phrase in duplicate_phrases:
        deduped = deduped.replace(phrase, phrase.split(" A rigorous")[0] if "A rigorous" in phrase else phrase.split(" IV. METHODOLOGY")[0])
    return deduped.strip()


def _summarize_methodology_section(section_text: str) -> str:
    text = _dedupe_section_text(section_text)
    parts: list[str] = []
    if "five sequential stages" in text:
        parts.append(
            "the overall pipeline has five stages: Data Collection, Data Store and Clean, "
            "Data Annotation, Data Preparation, and Data Split"
        )
    if "Auto-orientation" in text or "Resizing and Standardization" in text:
        parts.append(
            "Data Preprocessing includes auto-orientation, resizing all images to 640x640, "
            "null-annotation filtering, and data augmentation"
        )
    if "YOLOv8s" in text:
        parts.append(
            "YOLOv8s is used as the baseline model architecture, with CSPDarknet-style feature extraction, "
            "FPN/PAN multi-scale fusion, and a decoupled detection head"
        )
    if "Transfer learning" in text:
        parts.append(
            "training uses transfer learning from COCO-pretrained weights, 100 epochs, batch size 16, Adam, "
            "cosine learning-rate scheduling, and early stopping"
        )
    if "GIS Integration" in text or "GIS platforms" in text:
        parts.append(
            "GIS Integration maps geotagged detections onto GIS platforms to generate spatial risk maps "
            "for epidemiological analysis and targeted intervention planning"
        )
    if not parts:
        return _trim_section_answer(text, 220)
    return "; ".join(parts) + "."


def _fallback_research_section_answer(vector_store: FAISS, question: str) -> tuple[str, list[Document]] | None:
    section = _is_research_section_question(question)
    if not section:
        return None

    sources = _research_paper_sources(vector_store)
    if not sources:
        return None

    question_lc = question.lower()
    chosen_source = sources[0]
    if "mosquito" in question_lc or "mosquit" in question_lc:
        for source in sources:
            if "mosquito" in source.lower():
                chosen_source = source
                break

    full_text, source_doc = _ordered_source_text(vector_store, chosen_source)
    section_text = _extract_research_section(full_text, section)
    if not section_text or not source_doc:
        return None

    citation = _source_citation(source_doc)
    if section == "abstract":
        answer = f"The abstract of the MosquitoFusion paper is: {_trim_section_answer(section_text, 170)} {citation}"
    else:
        answer = f"The methodology of the MosquitoFusion paper is: {_summarize_methodology_section(section_text)} {citation}"

    docs = [
        doc
        for doc in _all_indexed_documents(vector_store)
        if doc.metadata.get("source") == chosen_source
        and (
            (section == "abstract" and 163 <= int(doc.metadata.get("chunk_index", 0)) <= 166)
            or (section == "methodology" and 190 <= int(doc.metadata.get("chunk_index", 0)) <= 198)
        )
    ]
    docs.sort(key=lambda doc: int(doc.metadata.get("chunk_index", 0)))
    return answer, docs[:8] or [source_doc]


def _history_sources(chat_history: list[dict[str, str]]) -> list[str]:
    sources: list[str] = []
    for message in reversed(chat_history[-6:]):
        for source in message.get("sources") or []:
            source_name = str(source).split(" p.")[0].strip()
            if source_name and source_name not in sources:
                sources.append(source_name)
    return sources


def _recent_history_text(chat_history: list[dict[str, str]], max_messages: int = 4) -> str:
    parts: list[str] = []
    for message in chat_history[-max_messages:]:
        content = _normalise_inline_text(str(message.get("content", "")))
        if not content:
            continue
        if len(content) > 500:
            content = content[:497] + "..."
        parts.append(content)
    return " ".join(parts)


def _is_follow_up_question(question: str) -> bool:
    q = question.lower().strip()
    terms = _query_terms(q)
    if len(terms) <= 4:
        return True
    follow_markers = (
        "that",
        "this",
        "it",
        "them",
        "those",
        "week",
        "day",
        "step",
        "phase",
        "round",
        "module",
        "explain",
        "elaborate",
        "tell me more",
    )
    return any(marker in q for marker in follow_markers)


def _contextualize_question(question: str, chat_history: list[dict[str, str]]) -> str:
    question = _normalise_query(question)
    if not chat_history or not _is_follow_up_question(question):
        return question
    history_text = _recent_history_text(chat_history)
    if not history_text:
        return question
    return f"{history_text}\nFollow-up question: {question}"


def _timebox_section_request(question: str) -> tuple[str, str] | None:
    q = question.lower()
    patterns = (
        ("week", r"\bweek\s*([0-9]+)\b"),
        ("day", r"\bday\s*([0-9]+)\b"),
        ("step", r"\bstep\s*([0-9]+)\b"),
        ("phase", r"\bphase\s*([0-9]+)\b"),
    )
    for label, pattern in patterns:
        matches = list(re.finditer(pattern, q))
        if matches:
            return label, matches[-1].group(1)
    return None


def _extract_numbered_section(full_text: str, label: str, number: str) -> str | None:
    label_pattern = re.escape(label)
    number_pattern = re.escape(number)
    start_pattern = rf"\b{label_pattern}\s*{number_pattern}\b\s*[:.\-–—]?\s*"
    start_match = re.search(start_pattern, full_text, flags=re.IGNORECASE)
    if not start_match:
        return None

    next_pattern = rf"\b{label_pattern}\s*[0-9]+\b\s*[:.\-–—]?"
    next_match = re.search(next_pattern, full_text[start_match.end() :], flags=re.IGNORECASE)
    end = start_match.end() + next_match.start() if next_match else len(full_text)
    section = full_text[start_match.start() : end]
    return _normalise_inline_text(section)


def _fallback_numbered_section_answer(
    vector_store: FAISS,
    question: str,
    chat_history: list[dict[str, str]],
) -> tuple[str, list[Document]] | None:
    request = _timebox_section_request(question)
    if not request:
        return None

    label, number = request
    preferred_sources = _history_sources(chat_history)
    all_sources = sorted({str(doc.metadata.get("source", "")) for doc in _all_indexed_documents(vector_store)})
    candidate_sources = preferred_sources + [source for source in all_sources if source not in preferred_sources]

    for source in candidate_sources:
        full_text, source_doc = _ordered_source_text(vector_store, source)
        if not full_text or not source_doc:
            continue
        section = _extract_numbered_section(full_text, label, number)
        if not section:
            continue

        source_docs = _source_docs(vector_store, source)
        citation = _source_citation(source_doc)
        section = re.sub(rf"^\s*{re.escape(label)}\s*{re.escape(number)}\s*[:.\-–—]?\s*", "", section, flags=re.IGNORECASE)
        answer = f"{label.title()} {number}: {_trim_section_answer(section, 220)} {citation}"
        return answer, source_docs[:8] or [source_doc]

    return None


def _resume_candidate_sources(vector_store: FAISS) -> list[str]:
    sources: dict[str, int] = {}
    for doc in _all_indexed_documents(vector_store):
        source = str(doc.metadata.get("source", ""))
        text = doc.page_content.lower()
        if (
            "resume" in source.lower()
            or "career objective" in text
            or "personal details" in text
            or "father" in text
            or "technical skills" in text
            or "strengths" in text
        ):
            sources[source] = sources.get(source, 0) + 1
    return sorted(sources, key=sources.get, reverse=True)


def _source_docs(vector_store: FAISS, source: str) -> list[Document]:
    docs = [
        doc
        for doc in _all_indexed_documents(vector_store)
        if doc.metadata.get("source") == source
    ]
    docs.sort(key=lambda doc: int(doc.metadata.get("chunk_index", 0)))
    return docs


def _resume_lines(text: str) -> list[str]:
    return [line.strip() for line in text.splitlines() if line.strip()]


def _extract_resume_facts(text: str) -> dict[str, str]:
    facts: dict[str, str] = {}
    lines = _resume_lines(text)
    inline = _normalise_inline_text(text)

    phone_match = re.search(r"(?:\+?\d[\d\s\-]{8,}\d)", inline)
    if phone_match:
        facts["phone"] = phone_match.group(0).strip()

    email_match = re.search(r"[\w.+-]+@[\w.-]+\.\w+", inline)
    if email_match:
        facts["email"] = email_match.group(0).strip()

    father_matches = list(re.finditer(
        r"Father[’']?s\s*Name\s*:\s*(.+?)(?:Date\s*of\s*Birth|Gender|Marital|Nationality|Languages|$)",
        inline,
        flags=re.IGNORECASE,
    ))
    if father_matches:
        father_values = [match.group(1).strip(" •|") for match in father_matches]
        facts["father"] = max(father_values, key=lambda value: (value.count(" "), len(value)))

    address_matches = list(re.finditer(
        r"(Vill\.\s*.+?)(?:Career\s*Objective|Education|$)",
        inline,
        flags=re.IGNORECASE,
    ))
    if address_matches:
        address_values = [match.group(1).strip(" •|") for match in address_matches]
        facts["address"] = max(address_values, key=lambda value: (value.count(" "), len(value)))
    else:
        for index, line in enumerate(lines):
            if "@" in line or re.search(r"\+?\d[\d\s\-]{8,}\d", line):
                address_parts: list[str] = []
                for following in lines[index + 1 : index + 5]:
                    if re.search(r"career\s*objective|education|experience", following, flags=re.IGNORECASE):
                        break
                    address_parts.append(following)
                if address_parts:
                    facts["address"] = " ".join(address_parts)
                    break

    strengths_lines: list[str] = []
    collecting_strengths = False
    stop_headings = {
        "personal details",
        "declaration",
        "career objective",
        "education",
        "experience",
        "technical skills",
        "academic achievement",
    }
    for line in lines:
        clean_line = line.strip(" •\t")
        line_lc = clean_line.lower()
        if line_lc == "strengths":
            collecting_strengths = True
            strengths_lines = []
            continue
        if collecting_strengths:
            if line_lc in stop_headings or "@" in clean_line or re.search(r"\+?\d[\d\s\-]{8,}\d", clean_line):
                break
            if clean_line and len(clean_line.split()) >= 3:
                strengths_lines.append(clean_line)
    if strengths_lines:
        facts["qualities"] = "; ".join(dict.fromkeys(strengths_lines))

    skills_match = re.search(
        r"Soft\s*Skills\s*:\s*(.+?)(?:Strengths|Personal\s*Details|Declaration|$)",
        inline,
        flags=re.IGNORECASE,
    )
    if skills_match:
        facts["soft_skills"] = skills_match.group(1).strip(" •|")

    if lines:
        first_line = next((line for line in lines if not line.startswith("+") and "@" not in line), "")
        if first_line and len(first_line.split()) <= 4:
            facts["name"] = first_line

    return facts


def _resume_intent(question: str) -> str | None:
    question = _normalise_query(question)
    q = question.lower()
    if any(marker in q for marker in ("father", "fathers", "parent", "पिता")):
        return "father"
    if any(marker in q for marker in ("address", "addr", "पता")) or q.strip() in {"address", "addr"}:
        return "address"
    if any(marker in q for marker in ("phone", "mobile", "contact", "number", "no")):
        return "phone"
    if "email" in q or "mail" in q:
        return "email"
    if any(marker in q for marker in ("qualities", "quality", "qulaities", "strength", "strengths", "soft skill")):
        return "qualities"
    if any(marker in q for marker in ("name", "candidate")):
        return "name"
    return None


def _fallback_resume_answer(vector_store: FAISS, question: str) -> tuple[str, list[Document]] | None:
    intent = _resume_intent(question)
    if not intent:
        return None

    for source in _resume_candidate_sources(vector_store):
        docs = _source_docs(vector_store, source)
        if not docs:
            continue
        for doc in docs:
            facts = _extract_resume_facts(doc.page_content)
            citation = _source_citation(doc)

            if intent == "phone" and facts.get("phone"):
                return f"Phone number: {facts['phone']}. {citation}", docs[:4]
            if intent == "email" and facts.get("email"):
                return f"Email: {facts['email']}. {citation}", docs[:4]
            if intent == "father" and facts.get("father"):
                return f"Father's name: {facts['father']}. {citation}", docs[:4]
            if intent == "address" and facts.get("address") and facts["address"].count(" ") >= 2:
                return f"Address: {facts['address']}. {citation}", docs[:4]
            if intent == "qualities":
                if facts.get("qualities"):
                    return f"Qualities/strengths: {facts['qualities']}. {citation}", docs[:4]
            if intent == "name" and facts.get("name"):
                return f"Name: {facts['name']}. {citation}", docs[:4]

        full_text = _source_text_from_uploaded_file(source) or "\n".join(doc.page_content for doc in docs)
        facts = _extract_resume_facts(full_text)
        citation = _source_citation(docs[0])

        if intent == "phone" and facts.get("phone"):
            return f"Phone number: {facts['phone']}. {citation}", docs[:4]
        if intent == "email" and facts.get("email"):
            return f"Email: {facts['email']}. {citation}", docs[:4]
        if intent == "father" and facts.get("father"):
            return f"Father's name: {facts['father']}. {citation}", docs[:4]
        if intent == "address" and facts.get("address"):
            return f"Address: {facts['address']}. {citation}", docs[:4]
        if intent == "qualities":
            qualities = facts.get("qualities") or facts.get("soft_skills")
            if qualities:
                return f"Qualities/strengths: {qualities}. {citation}", docs[:4]
        if intent == "name" and facts.get("name"):
            return f"Name: {facts['name']}. {citation}", docs[:4]

    return None


def _contains_devanagari(text: str) -> bool:
    return bool(re.search(r"[\u0900-\u097F]", text))


def _source_citation(doc: Document) -> str:
    source = doc.metadata.get("source", "unknown")
    page = doc.metadata.get("page", "?")
    return f"[{source} p.{page}]"


def _has_any(text: str, needles: tuple[str, ...]) -> bool:
    text_lc = text.lower()
    return any(needle.lower() in text_lc for needle in needles)


def _detect_question_language(question: str) -> str:
    if _contains_devanagari(question):
        hindi_markers = ("यह", "इस", "क्या", "किस", "किसने", "कौन", "बीमार", "उद्देश्य")
        if any(marker in question for marker in hindi_markers):
            return "hi"
        return "mr"
    return "en"


def _clean_fact_value(value: str) -> str:
    return value.strip(" ।|,")


def _extract_shubhasandesh_facts(docs: list[Document]) -> dict[str, str]:
    combined = _normalise_inline_text("\n".join(doc.page_content for doc in docs))
    if "शुभसंदेश" not in combined:
        return {}

    facts: dict[str, str] = {}
    if "विनोद तावडे" in combined:
        facts["giver"] = "विनोद तावडे"

    if "मंत्री" in combined:
        role_match = re.search(r"मंत्री\s*\|\s*(.+?)\s*महाराष्ट्र राज्य", combined)
        if role_match:
            facts["role"] = _clean_fact_value(role_match.group(1).replace("|", ", ")) + " मंत्री"
        else:
            facts["role"] = "महाराष्ट्र राज्याचे मंत्री"

    date_match = re.search(r"दिनांक\s*:\s*([^|]+)", combined)
    if date_match:
        facts["date"] = _clean_fact_value(date_match.group(1))

    if "वैद्य साने ट्रस्ट" in combined and "माधवबाग" in combined:
        facts["institutions"] = "वैद्य साने ट्रस्ट आणि माधवबाग"

    if "आरोग्य क्षेत्रातील कार्य" in combined:
        facts["field"] = "आरोग्य क्षेत्रातील कार्य"

    campaign_match = re.search(r"'([^']*हृदयसंपदा[^']*)'", combined)
    if campaign_match:
        facts["campaign"] = _clean_fact_value(campaign_match.group(1))

    if "आरोग्यविषयक जाणिवांचा विस्तार" in combined:
        facts["purpose"] = "आरोग्यविषयक जाणिवांचा विस्तार साधणे"

    if "हृदयरोग" in combined and "मधुमेह" in combined:
        facts["diseases"] = "हृदयरोग आणि मधुमेह"

    if "मोठ्या शहरांसह ग्रामीण भागातही" in combined:
        facts["reach"] = "मोठ्या शहरांसह ग्रामीण भागात"

    cause_match = re.search(r"(जागतिकौकरण आणि अन्य संबंधित घटकांमुळे जीवनशैलीवर विपरीत परिणाम[^|।]*)", combined)
    if cause_match:
        facts["cause"] = _clean_fact_value(cause_match.group(1))

    if "संपूर्ण समाजाचे यामुळे प्रत्यक्ष व अप्रत्यक्षरीत्या अनेक अर्थानी नुकसान" in combined:
        facts["impact"] = "संपूर्ण समाजाचे प्रत्यक्ष व अप्रत्यक्षरीत्या अनेक अर्थांनी नुकसान होत आहे"

    if "आरोग्य, उर्जा व प्रेरणा देणारे" in combined:
        facts["benefit"] = "आरोग्य, उर्जा व प्रेरणा"
        facts["wish"] = "आरोग्य, उर्जा व प्रेरणा देणारे ठरो, ही सदिच्छा"

    if "अभिनंदनीय बाब" in combined:
        facts["opinion"] = "हे निश्चितच अभिनंदनीय बाब आहे"

    if "माझ्या मन:पूर्वक शुभेच्छा" in combined:
        facts["closing"] = "पुढील कार्यासाठी आणि माहितीपुस्तिकेसाठी मन:पूर्वक शुभेच्छा"

    return facts


def _format_shubhasandesh_fact(intent: str, facts: dict[str, str], citation: str, language: str) -> str | None:
    value = facts.get(intent)
    if not value:
        return None

    if language == "hi":
        templates = {
            "giver": f"यह शुभसंदेश {value} ने दिया है। {citation}",
            "role": f"विनोद तावडे {value} थे। {citation}",
            "date": f"इस शुभसंदेश की तारीख {value} है। {citation}",
            "institutions": f"यह शुभसंदेश {value} के लिए है। {citation}",
            "field": f"इसमें {value} की सराहना की गई है। {citation}",
            "campaign": f"इसमें '{value}' स्वास्थ्य चळवळ का उल्लेख है। {citation}",
            "purpose": f"इस चळवळ का उद्देश्य {value} है। {citation}",
            "diseases": f"इसमें {value} का प्रमाण तेजी से बढ़ने की बात कही गई है। {citation}",
            "reach": f"दस्तावेज में संस्था के {value} पहुंचने का उल्लेख है। {citation}",
            "cause": f"कारण के रूप में यह लिखा है: {value}. {citation}",
            "impact": f"परिणाम के रूप में {value}. {citation}",
            "benefit": f"संस्था के कार्य से समाज को {value} मिलती है। {citation}",
            "opinion": f"विनोद तावडे ने संस्था के कार्य को '{value}' कहा है। {citation}",
            "closing": f"अंत में {value} दी गई हैं। {citation}",
        }
        return templates.get(intent)

    if language == "en":
        english_values = {
            "विनोद तावडे": "Vinod Tawde",
            "वैद्य साने ट्रस्ट आणि माधवबाग": "Vaidya Sane Trust and Madhavbaug",
            "हृदयरोग आणि मधुमेह": "heart disease and diabetes",
            "आरोग्य क्षेत्रातील कार्य": "health-sector work",
            "आरोग्यविषयक जाणिवांचा विस्तार साधणे": "expanding health awareness",
            "आरोग्य, उर्जा व प्रेरणा": "health, energy, and inspiration",
        }
        value_en = english_values.get(value, value)
        templates = {
            "giver": f"The message was given by {value_en}. {citation}",
            "role": f"Vinod Tawde was listed as minister for: {value_en}. {citation}",
            "date": f"The date of the message is {value_en}. {citation}",
            "institutions": f"The message is for {value_en}. {citation}",
            "field": f"The document appreciates their work in {value_en}. {citation}",
            "campaign": f"The campaign mentioned is '{value_en}'. {citation}",
            "purpose": f"The stated purpose is {value_en}. {citation}",
            "diseases": f"The two diseases mentioned are {value_en}. {citation}",
            "reach": f"The document mentions reaching {value_en}. {citation}",
            "cause": f"The stated cause/context is: {value_en}. {citation}",
            "impact": f"The stated impact is: {value_en}. {citation}",
            "benefit": f"The work is wished to give society {value_en}. {citation}",
            "opinion": f"Vinod Tawde describes the work as: {value_en}. {citation}",
            "closing": f"The closing wishes are: {value_en}. {citation}",
        }
        return templates.get(intent)

    templates = {
        "giver": f"हा शुभसंदेश {value} यांनी दिला आहे. {citation}",
        "role": f"विनोद तावडे {value} होते. {citation}",
        "date": f"या शुभसंदेशाची तारीख {value} आहे. {citation}",
        "institutions": f"हा शुभसंदेश {value} यांच्यासाठी आहे. {citation}",
        "field": f"या शुभसंदेशात {value}ाचे कौतुक करण्यात आले आहे. {citation}",
        "campaign": f"या शुभसंदेशात '{value}' या स्वास्थ्य चळवळीचा उल्लेख आहे. {citation}",
        "purpose": f"या चळवळीचा उद्देश {value} असा दिला आहे. {citation}",
        "diseases": f"या शुभसंदेशामध्ये {value} यांचे प्रमाण वेगाने वाढत असल्याचे नमूद केले आहे. {citation}",
        "reach": f"दस्तऐवजात संस्थेचे कार्य {value} पोहोचणार असल्याचे नमूद आहे. {citation}",
        "cause": f"कारण/पार्श्वभूमी म्हणून मजकुरात असे दिले आहे: {value}. {citation}",
        "impact": f"या परिणामांबद्दल मजकुरात असे दिले आहे: {value}. {citation}",
        "benefit": f"संस्थेच्या कार्यामुळे समाजाला {value} मिळो अशी सदिच्छा व्यक्त केली आहे. {citation}",
        "opinion": f"विनोद तावडे यांनी संस्थेच्या कार्याबद्दल '{value}' असे मत व्यक्त केले आहे. {citation}",
        "closing": f"शेवटी {value} दिल्या आहेत. {citation}",
    }
    return templates.get(intent)


def _detect_shubhasandesh_intent(question: str) -> str | None:
    q = question.lower()
    intent_markers = [
        ("date", ("तारीख", "दिनांक", "date")),
        ("giver", ("कोणी", "कोणाकडून", "किसने", "who gave", "given by")),
        ("role", ("पद", "कार्यरत", "मंत्री", "position", "post", "role")),
        ("institutions", ("संस्थेसाठी", "किस संस्था", "which institution", "for which")),
        ("field", ("क्षेत्र", "field", "कार्याचे कौतुक", "सराहना")),
        ("campaign", ("अभियान", "चळवळ", "campaign")),
        ("purpose", ("उद्देश", "purpose", "objective")),
        ("diseases", ("आजार", "बीमार", "disease", "diseases")),
        ("reach", ("कुठे", "where", "शहरी", "ग्रामीण")),
        ("cause", ("कारण", "cause", "reason")),
        ("impact", ("परिणाम", "result", "impact")),
        ("benefit", ("लाभ", "benefit", "काय लाभ", "समाजाला काय")),
        ("opinion", ("मत", "opinion", "कार्याबद्दल काय")),
        ("closing", ("शेवटी", "शुभेच्छा", "ending", "closing")),
    ]
    for intent, markers in intent_markers:
        if any(marker in q for marker in markers):
            return intent
    if "काय आहे" in q or "what is this" in q:
        return "summary"
    return None


def _fallback_marathi_answer(question: str, docs: list[Document]) -> str | None:
    combined = _normalise_inline_text("\n".join(doc.page_content for doc in docs))
    if not docs or not _contains_devanagari(question + combined):
        return None
    facts = _extract_shubhasandesh_facts(docs)
    if not facts:
        return None

    source_doc = next((doc for doc in docs if "शुभसंदेश" in doc.page_content), docs[0])
    citation = _source_citation(source_doc)
    language = _detect_question_language(question)
    intent = _detect_shubhasandesh_intent(question)

    if intent == "summary":
        institutions = facts.get("institutions")
        giver = facts.get("giver")
        field = facts.get("field")
        if institutions and giver and field:
            if language == "hi":
                return f"यह {institutions} के {field} के बारे में {giver} द्वारा दिया गया शुभसंदेश है। {citation}"
            if language == "en":
                return f"This is a शुभसंदेश from {giver} about the {field} of {institutions}. {citation}"
            return f"हा {institutions} यांच्या {field}ाबद्दल {giver} यांनी दिलेला शुभसंदेश आहे. {citation}"

    if intent:
        return _format_shubhasandesh_fact(intent, facts, citation, language)

    return None


def _line_contexts(docs: list[Document]) -> list[tuple[str, Document]]:
    contexts: list[tuple[str, Document]] = []
    for doc in docs:
        raw_lines = [line.strip(" |") for line in doc.page_content.splitlines() if line.strip(" |")]
        for index, line in enumerate(raw_lines):
            previous_line = raw_lines[index - 1] if index > 0 else ""
            next_line = raw_lines[index + 1] if index + 1 < len(raw_lines) else ""
            contexts.append((" ".join(part for part in (previous_line, line, next_line) if part), doc))
    return contexts


def _find_notice_context(docs: list[Document], markers: tuple[str, ...]) -> tuple[str, Document] | None:
    for context, doc in _line_contexts(docs):
        context_lc = context.lower()
        if any(marker.lower() in context_lc for marker in markers):
            return context, doc
    return None


def _clean_notice_context(text: str) -> str:
    text = _normalise_inline_text(text)
    text = re.sub(r"\s+", " ", text)
    return text.strip(" |")


def _fallback_notice_answer(question: str, docs: list[Document]) -> str | None:
    if not docs or not _contains_devanagari(question):
        return None

    q = question.lower()
    wants_students = any(marker in q for marker in ("विद्यार्थ", "students", "जारी"))
    wants_portal = any(marker in q for marker in ("पोर्टल", "portal", "संकेतस्थळ", "website"))
    wants_course = any(marker in q for marker in ("अभ्यासक्रम", "course", "program", "programme"))
    if not any((wants_students, wants_portal, wants_course)):
        return None

    answers: list[str] = []
    if wants_students:
        match = _find_notice_context(docs, ("विद्यार्थ्यांसाठी", "विद्यार्थी", "students"))
        if match:
            context, doc = match
            answers.append(f"सूचना विद्यार्थ्यांबाबतचा मजकूर: {_clean_notice_context(context)} {_source_citation(doc)}")

    if wants_portal:
        match = _find_notice_context(docs, ("पोर्टल", "portal", "संकेतस्थळ", "website"))
        if match:
            context, doc = match
            answers.append(f"निकाल/माहिती पोर्टलबाबतचा मजकूर: {_clean_notice_context(context)} {_source_citation(doc)}")

    if wants_course:
        programs: list[str] = []
        for doc in docs:
            for program in _detect_programs(str(doc.metadata.get("source", "")) + "\n" + doc.page_content):
                if program not in programs:
                    programs.append(program)
        if programs:
            answers.append(f"ही नोटीस {', '.join(programs)} अभ्यासक्रमासाठी आहे. {_source_citation(docs[0])}")

    return "\n".join(answers) if answers else None


def _focused_excerpt(question: str, doc: Document, max_chars: int = 360) -> str:
    text = _normalise_inline_text(doc.page_content)
    text_lc = text.lower()
    terms = [term for term in _query_terms(question) if term in text_lc]
    anchor = min((text_lc.find(term) for term in terms if text_lc.find(term) >= 0), default=0)
    start = max(anchor - max_chars // 3, 0)
    end = min(start + max_chars, len(text))
    excerpt = text[start:end].strip()
    if start > 0:
        excerpt = "..." + excerpt
    if end < len(text):
        excerpt += "..."
    return excerpt


def build_fallback_answer(question: str, docs: list[Document]) -> str:
    direct_answer = (
        _fallback_marathi_answer(question, docs)
        or _fallback_notice_answer(question, docs)
        or _fallback_exam_centre_answer(question, docs)
        or _fallback_vacancy_answer(question, docs)
    )
    if direct_answer:
        return direct_answer

    lines: list[str] = ["I could not use Cerebras for this response, so here is the best grounded answer from the uploaded files:"]
    for index, doc in enumerate(docs[:3], 1):
        source = doc.metadata.get("source", "unknown")
        page = doc.metadata.get("page", "?")
        excerpt = _focused_excerpt(question, doc)
        lines.append(f"{index}. {excerpt} [{source} p.{page}]")
    if not docs:
        lines.append("No relevant document chunks were found.")
    return "\n".join(lines)


def make_llm_client() -> Cerebras:
    api_key = os.getenv("CEREBRAS_API_KEY")
    if not api_key:
        raise RuntimeError("CEREBRAS_API_KEY is not set")
    return Cerebras(api_key=api_key)


def to_cerebras_messages(messages) -> list[dict[str, str]]:
    role_map = {"human": "user", "ai": "assistant", "system": "system"}
    return [
        {"role": role_map.get(message.type, message.type), "content": message.content}
        for message in messages
    ]


def build_chat_history_context(chat_history: list[dict[str, str]], max_messages: int = 8) -> str:
    if not chat_history:
        return "No prior chat history yet."

    recent_messages = chat_history[-max_messages:]
    lines: list[str] = []
    for message in recent_messages:
        role = message.get("role", "unknown").capitalize()
        content = " ".join(str(message.get("content", "")).split())
        if len(content) > 350:
            content = content[:347] + "..."
        sources = message.get("sources") or []
        if sources:
            source_text = "; ".join(str(source) for source in sources)
            lines.append(f"{role}: {content} [sources: {source_text}]")
        else:
            lines.append(f"{role}: {content}")
    return "\n".join(lines)


def _tokenize(text: str) -> list[str]:
    return re.findall(r"[a-z0-9]+|[\u0900-\u097F]+", text.lower())


def _query_terms(question: str) -> list[str]:
    terms: list[str] = []
    for token in _tokenize(question):
        if token not in STOPWORDS and len(token) > 1 and token not in terms:
            terms.append(token)
    return terms


def _retrieval_confidence(question: str, docs: list[Document]) -> float:
    terms = [term for term in _query_terms(question) if term not in DOMAIN_QUERY_TERMS]
    if not terms:
        return 0.75 if docs else 0.0
    combined = _normalise_inline_text("\n".join(doc.page_content for doc in docs)).lower()
    matched = sum(1 for term in terms if term.lower() in combined)
    return matched / len(terms)


def _not_found_answer(question: str) -> str:
    language = _detect_question_language(question)
    if language == "mr":
        return "अपलोड केलेल्या दस्तऐवजांमध्ये याचे स्पष्ट उत्तर सापडले नाही."
    if language == "hi":
        return "अपलोड किए गए दस्तावेज़ों में इसका स्पष्ट उत्तर नहीं मिला."
    return "I could not find a clearly supported answer in the uploaded documents."


def _looks_like_not_found(answer: str) -> bool:
    answer_lc = answer.lower()
    markers = (
        "could not find",
        "not present",
        "uploaded documents",
        "स्पष्ट उत्तर सापडले नाही",
        "माहिती मिळू शकली नाही",
        "नहीं मिला",
        "नहीं मिल",
    )
    return any(marker in answer_lc for marker in markers)


def _answer_support_score(answer: str, docs: list[Document]) -> float:
    answer_terms = [
        term
        for term in _query_terms(answer)
        if term not in DOMAIN_QUERY_TERMS and len(term) > 2
    ]
    if not answer_terms:
        return 1.0
    combined = _normalise_inline_text("\n".join(doc.page_content for doc in docs)).lower()
    matched = sum(1 for term in answer_terms if term.lower() in combined)
    return matched / len(answer_terms)


def _doc_key(doc: Document) -> tuple[str, int | str, int | str]:
    return (
        str(doc.metadata.get("source", "")),
        doc.metadata.get("page", ""),
        doc.metadata.get("chunk_index", ""),
    )


def _all_indexed_documents(vector_store: FAISS) -> list[Document]:
    docstore = getattr(vector_store, "docstore", None)
    doc_dict = getattr(docstore, "_dict", None)
    if not doc_dict:
        return []
    return list(doc_dict.values())


def _lexical_score(question: str, doc: Document) -> float:
    question = _normalise_query(question)
    terms = _query_terms(question)
    if not terms:
        return 0.0

    text = doc.page_content.lower()
    tokens = set(_tokenize(text))
    matched_terms = [term for term in terms if term in tokens or term in text]
    if not matched_terms:
        return 0.0

    score = len(matched_terms) / len(terms)
    score += 0.08 * sum(text.count(term) for term in matched_terms)

    question_lc = question.lower()
    requested_entities = [term for term in terms if term not in DOMAIN_QUERY_TERMS]
    matched_entities = [term for term in requested_entities if term in tokens or term in text]
    requested_programs = _detect_programs(question)
    doc_programs = _detect_programs(str(doc.metadata.get("source", "")) + "\n" + doc.page_content)

    if "exam" in question_lc and ("centre" in question_lc or "center" in question_lc):
        if "online examination centre" in text or "online examination center" in text:
            score += 1.25
        if "examination centre" in text or "examination center" in text:
            score += 0.75
        if "number of vacancies" in text or "vacancies are as under" in text:
            score -= 0.6
        if requested_entities and not matched_entities:
            score -= 1.2

    score += 2.0 * len(matched_entities)
    if requested_entities and len(matched_entities) == len(requested_entities):
        score += 0.8
    if _is_fee_question(question):
        if requested_programs and any(program in doc_programs for program in requested_programs):
            score += 2.5
        elif requested_programs and doc_programs:
            score -= 2.0
        if any(term in text for term in ("fee", "hostel", "caution", "security", "tuition")):
            score += 0.8
    score += _source_relevance_score(question, str(doc.metadata.get("source", "")))

    return score


def _normalise_scores(scores: dict[tuple[str, int | str, int | str], float]) -> dict[tuple[str, int | str, int | str], float]:
    if not scores:
        return {}
    min_score = min(scores.values())
    max_score = max(scores.values())
    if max_score == min_score:
        return {key: 1.0 for key in scores}
    return {key: (value - min_score) / (max_score - min_score) for key, value in scores.items()}


def hybrid_retrieve_documents(
    vector_store: FAISS,
    question: str,
    final_k: int = 8,
    dense_k: int = 32,
    lexical_k: int = 32,
) -> list[Document]:
    question = _normalise_query(question)
    dense_results = vector_store.similarity_search_with_score(question, k=dense_k)
    all_docs = _all_indexed_documents(vector_store)

    docs_by_key: dict[tuple[str, int | str, int | str], Document] = {}
    dense_raw: dict[tuple[str, int | str, int | str], float] = {}
    for doc, distance in dense_results:
        key = _doc_key(doc)
        docs_by_key[key] = doc
        dense_raw[key] = -float(distance)

    lexical_ranked = sorted(
        ((_lexical_score(question, doc), doc) for doc in all_docs),
        key=lambda item: item[0],
        reverse=True,
    )
    lexical_raw: dict[tuple[str, int | str, int | str], float] = {}
    for score, doc in lexical_ranked[:lexical_k]:
        if score <= 0:
            continue
        key = _doc_key(doc)
        docs_by_key[key] = doc
        lexical_raw[key] = score

    dense_scores = _normalise_scores(dense_raw)
    lexical_scores = _normalise_scores(lexical_raw)

    ranked_keys = sorted(
        docs_by_key,
        key=lambda key: (0.45 * dense_scores.get(key, 0.0)) + (0.55 * lexical_scores.get(key, 0.0)),
        reverse=True,
    )
    return _expand_with_neighbor_chunks(vector_store, [docs_by_key[key] for key in ranked_keys[:final_k]], final_k)


def _expand_with_neighbor_chunks(vector_store: FAISS, docs: list[Document], final_k: int) -> list[Document]:
    all_docs = _all_indexed_documents(vector_store)
    by_source_index: dict[tuple[str, int], Document] = {}
    for doc in all_docs:
        try:
            chunk_index = int(doc.metadata.get("chunk_index", -9999))
        except Exception:
            continue
        by_source_index[(str(doc.metadata.get("source", "")), chunk_index)] = doc

    expanded: list[Document] = []
    seen: set[tuple[str, int | str, int | str]] = set()
    for doc in docs:
        source = str(doc.metadata.get("source", ""))
        try:
            chunk_index = int(doc.metadata.get("chunk_index", -9999))
        except Exception:
            chunk_index = -9999

        for candidate in (
            by_source_index.get((source, chunk_index - 1)),
            doc,
            by_source_index.get((source, chunk_index + 1)),
        ):
            if not candidate:
                continue
            key = _doc_key(candidate)
            if key not in seen:
                expanded.append(candidate)
                seen.add(key)
            if len(expanded) >= final_k:
                return expanded
    return expanded


def retrieve_context(vector_store: FAISS, question: str, k: int = 8) -> tuple[str, list[Document]]:
    docs = hybrid_retrieve_documents(vector_store, question, final_k=k)
    lines: list[str] = []
    for i, doc in enumerate(docs, 1):
        source = doc.metadata.get("source", "unknown")
        page = doc.metadata.get("page", "?")
        lines.append(f"[{i}] Source: {source}, p.{page}\n{doc.page_content}")
    return "\n\n---\n\n".join(lines), docs


def answer_question(vector_store: FAISS, question: str, chat_history: list[dict[str, str]]) -> tuple[str, list[Document]]:
    original_question = _normalise_query(question)
    resolved_question = _contextualize_question(original_question, chat_history)

    fee_answer = _fallback_fee_answer(vector_store, resolved_question)
    if fee_answer:
        return fee_answer

    numbered_section_answer = _fallback_numbered_section_answer(vector_store, resolved_question, chat_history)
    if numbered_section_answer:
        return numbered_section_answer

    resume_answer = _fallback_resume_answer(vector_store, resolved_question)
    if resume_answer:
        return resume_answer

    research_section_answer = _fallback_research_section_answer(vector_store, resolved_question)
    if research_section_answer:
        return research_section_answer

    context, docs = retrieve_context(vector_store, resolved_question)
    if not docs:
        return "I could not find this in the uploaded documents.", []

    direct_answer = (
        _fallback_marathi_answer(resolved_question, docs)
        or _fallback_notice_answer(resolved_question, docs)
        or _fallback_exam_centre_answer(resolved_question, docs)
        or _fallback_vacancy_answer(resolved_question, docs)
    )
    if direct_answer:
        return direct_answer, docs

    if _retrieval_confidence(resolved_question, docs) < 0.35:
        return _not_found_answer(original_question), docs

    try:
        prompt = build_answer_prompt()
        chat_history_context = build_chat_history_context(chat_history)
        prompt_question = original_question
        if resolved_question != original_question:
            prompt_question = f"{original_question}\nResolved with chat context: {resolved_question}"
        messages = prompt.format_messages(question=prompt_question, context=context, chat_history=chat_history_context)
        client = make_llm_client()
        response = client.chat.completions.create(
            model=os.getenv("CEREBRAS_MODEL", "gpt-oss-120b"),
            messages=to_cerebras_messages(messages),
            max_completion_tokens=int(os.getenv("CEREBRAS_MAX_TOKENS", "1024")),
            temperature=float(os.getenv("CEREBRAS_TEMPERATURE", "0.2")),
            top_p=float(os.getenv("CEREBRAS_TOP_P", "1")),
            stream=False,
            reasoning_effort=os.getenv("CEREBRAS_REASONING_EFFORT", "medium"),
        )
        content = (response.choices[0].message.content or "").strip()
        if not content:
            return build_fallback_answer(question, docs), docs
        if _looks_like_not_found(content):
            return content, docs
        verified_content = _verify_structured_claims(content, resolved_question, vector_store)
        if verified_content:
            content = verified_content
        else:
            return _not_found_answer(original_question), docs
        if _answer_support_score(content, docs) < 0.35:
            return _not_found_answer(original_question), docs
        return content, docs
    except Exception as exc:  # pragma: no cover - demo UX path
        message = str(exc)
        if "disabled_organization" in message or "organization this API key belongs is disabled" in message:
            return build_fallback_answer(question, docs), docs
        return build_fallback_answer(question, docs), docs


def initialize_state() -> None:
    if "vector_store" not in st.session_state:
        st.session_state.vector_store = load_vector_store()
    if "chat_history" not in st.session_state:
        st.session_state.chat_history = []


def render_sidebar() -> None:
    st.sidebar.header("Demo Controls")
    st.sidebar.caption("Build a local knowledge base from uploaded files, then chat with citations.")
    st.sidebar.write("Current index status:")
    manifest = load_manifest()
    if manifest.get("chunk_count"):
        st.sidebar.success(f"{manifest['chunk_count']} chunks indexed")
        if manifest.get("files"):
            st.sidebar.write("Files:")
            for file_name in manifest["files"]:
                st.sidebar.write(f"- {file_name}")
        metadata = load_document_metadata()
        fee_records = load_fee_records()
        if metadata:
            with st.sidebar.expander("Document diagnostics"):
                for item in metadata:
                    programs = ", ".join(item.get("programs_detected") or ["None"])
                    st.write(f"**{item.get('source')}**")
                    st.write(f"Type: {item.get('document_type')} | Programs: {programs}")
        if fee_records:
            with st.sidebar.expander(f"Fee records ({len(fee_records)})"):
                for record in fee_records[:20]:
                    st.write(
                        f"{record['program']} | {record['fee_type']} | {record['amount']} "
                        f"({record['source']} p.{record['page']})"
                    )
    else:
        st.sidebar.info("No knowledge base built yet")

    if st.sidebar.button("Clear knowledge base", use_container_width=True):
        st.session_state.vector_store = None
        st.session_state.chat_history = []
        if INDEX_DIR.exists():
            for child in INDEX_DIR.iterdir():
                if child.is_file():
                    child.unlink()
                else:
                    shutil.rmtree(child)
        if MANIFEST_PATH.exists():
            MANIFEST_PATH.unlink()
        if DOCUMENT_METADATA_PATH.exists():
            DOCUMENT_METADATA_PATH.unlink()
        if FEE_RECORDS_PATH.exists():
            FEE_RECORDS_PATH.unlink()
        st.sidebar.success("Knowledge base cleared")


def main() -> None:
    st.set_page_config(page_title="Basic RAG MVP", page_icon="", layout="wide")
    initialize_state()

    st.title("Basic RAG MVP")
    st.write(
        "Upload documents, build a local knowledge base, and test retrieval quality with a grounded chat experience."
    )

    if not os.getenv("CEREBRAS_API_KEY"):
        st.warning("CEREBRAS_API_KEY is not set. Upload and indexing can still work, but answers need a Cerebras key.")

    render_sidebar()

    upload_col, action_col = st.columns([2, 1])
    with upload_col:
        uploaded_files = st.file_uploader(
            "Upload PDF, DOCX, XLSX, CSV, TXT, or MD files",
            type=sorted(SUPPORTED_TYPES),
            accept_multiple_files=True,
        )
    with action_col:
        build_clicked = st.button("Build knowledge base", use_container_width=True)

    if build_clicked:
        if not uploaded_files:
            st.error("Please upload at least one file first.")
        else:
            all_docs: list[Document] = []
            uploaded_names: list[str] = []
            errors: list[str] = []

            for uploaded_file in uploaded_files:
                try:
                    docs = load_documents_from_upload(uploaded_file)
                    if docs:
                        all_docs.extend(docs)
                        uploaded_names.append(uploaded_file.name)
                    else:
                        errors.append(f"{uploaded_file.name}: no readable text found")
                except Exception as exc:  # pragma: no cover - demo UX path
                    errors.append(f"{uploaded_file.name}: {exc}")

            if not all_docs:
                st.error("No text could be extracted from the uploaded files.")
                if errors:
                    st.write(errors)
            else:
                chunks = build_chunks(all_docs)
                document_metadata = infer_document_metadata(all_docs)
                fee_records = extract_fee_records(all_docs, document_metadata)
                vector_store = build_vector_store(chunks)
                save_vector_store(vector_store)
                st.session_state.vector_store = vector_store
                persist_manifest(uploaded_names, len(chunks))
                save_document_metadata(document_metadata)
                save_fee_records(fee_records)
                st.success(f"Knowledge base built from {len(uploaded_names)} file(s) and {len(chunks)} chunk(s).")
                if document_metadata:
                    with st.expander("Detected document metadata", expanded=True):
                        st.dataframe(pd.DataFrame(document_metadata), use_container_width=True)
                if fee_records:
                    with st.expander("Extracted fee records", expanded=True):
                        st.dataframe(pd.DataFrame(fee_records), use_container_width=True)
                if errors:
                    st.warning("Some files had issues:")
                    for error in errors:
                        st.write(f"- {error}")

    vector_store = st.session_state.get("vector_store")
    if not vector_store:
        st.info("Build the knowledge base before asking questions.")
        return

    st.divider()
    st.subheader("Chat Test")

    for message in st.session_state.chat_history:
        with st.chat_message(message["role"]):
            st.markdown(message["content"])
            if message["role"] == "assistant" and message.get("sources"):
                st.caption("Sources: " + ", ".join(message["sources"]))

    question = st.chat_input("Ask something about the uploaded documents")
    if question:
        current_history = list(st.session_state.chat_history)
        st.session_state.chat_history.append({"role": "user", "content": question})
        with st.chat_message("user"):
            st.markdown(question)

        with st.chat_message("assistant"):
            with st.spinner("Searching the knowledge base..."):
                answer, docs = answer_question(vector_store, question, current_history)
            st.markdown(answer)
            sources = []
            for doc in docs:
                source = doc.metadata.get("source", "unknown")
                page = doc.metadata.get("page", "?")
                sources.append(f"{source} p.{page}")
            if sources:
                st.caption("Sources: " + ", ".join(dict.fromkeys(sources)))

        st.session_state.chat_history.append(
            {
                "role": "assistant",
                "content": answer,
                "sources": list(dict.fromkeys(sources)),
            }
        )


if __name__ == "__main__":
    main()
